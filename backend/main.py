import warnings

warnings.filterwarnings(
    "ignore", 
    message="'pin_memory' argument is set as true but no accelerator is found"
)

import os
import time
import signal
import uuid
import tempfile
import asyncio
from datetime import datetime, timedelta
from typing import Dict, Optional, List, Set,Any,Tuple
from concurrent.futures import ThreadPoolExecutor
import hashlib
import subprocess
# Fix for OpenMP runtime conflict - must be at the very top
os.environ['KMP_DUPLICATE_LIB_OK'] = 'TRUE'

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Request, BackgroundTasks
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import mimetypes

from sqlalchemy.ext.asyncio import AsyncSession
from contextlib import asynccontextmanager

from database.database import engine, AsyncSessionLocal
from database.models import Base, User
from database.crud import create_user, authenticate_user, get_user_by_username

import httpx
import json

# RAG imports
import faiss
import numpy as np
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
import PyPDF2
from docx import Document as DocxDocument
import io
import pdfplumber
from docx.table import Table as DocxTable

from transformers import BlipProcessor, BlipForConditionalGeneration
from PIL import Image
import random
import easyocr
import fitz  # PyMuPDF

# Global variables
embedding_model = 'nomic-embed-text'  # Default Ollama embedding model
thread_pool = ThreadPoolExecutor(max_workers=4)
MODEL_PATH='http://localhost:'
EMBEDDING_PATH='http://localhost:8083'

AVAILABLE_PORT = 8081

caption_processor = BlipProcessor.from_pretrained("./blip_model",use_fast=False)
caption_model = BlipForConditionalGeneration.from_pretrained("./blip_model")

ocr_reader = easyocr.Reader(
            ['en'], 
            model_storage_directory="./models",
            download_enabled=True  
        )

# IMPROVED similarity threshold and retrieval parameters
SIMILARITY_THRESHOLD = 0.75  # Increased threshold for better relevance
MAX_CONTEXT_LENGTH = 3000    # Reduced to prevent overwhelming the model
MIN_CHUNK_OVERLAP = 100      # Increased overlap for better context continuity
CHUNK_SIZE = 800             # Smaller chunks for better precision
TOP_K_RESULTS = 3            # Fewer but more relevant results

# In-memory session storage
class RAGSession:
    
    def __init__(self, session_id: str):
        self.session_id = session_id
        self.created_at = datetime.now()
        self.last_accessed = datetime.now()
        self.document_name: Optional[str] = None
        self.chunks: List[Document] = []
        self.embeddings: Optional[np.ndarray] = None
        self.vector_store: Optional[faiss.IndexFlatIP] = None
        self.chunk_texts: List[str] = []
        self.tables: List[Dict[str, Any]] = []
        self.table_chunks: List[Document] = []
        self.document_metadata: Dict[str, Any] = {}

    def update_access_time(self):
        self.last_accessed = datetime.now()

    def has_document(self) -> bool:
        return self.document_name is not None

    def clear_document_data(self):
        self.document_name = None
        self.chunks = []
        self.embeddings = None
        self.vector_store = None
        self.chunk_texts = []
        self.tables = []
        self.table_chunks = []
        self.document_metadata = {}

# Global session storage
rag_sessions: Dict[str, RAGSession] = {}
SESSION_TIMEOUT_HOURS = 2

async def cleanup_expired_sessions():
    """Remove expired sessions"""
    current_time = datetime.now()
    expired_sessions = []
    
    for session_id, session in rag_sessions.items():
        if current_time - session.last_accessed > timedelta(hours=SESSION_TIMEOUT_HOURS):
            expired_sessions.append(session_id)
    
    for session_id in expired_sessions:
        del rag_sessions[session_id]
        print(f"Expired session removed: {session_id}")

@asynccontextmanager
async def lifespan(app: FastAPI):
    # Create database tables at startup
    try:
        async with engine.begin() as conn:
            await conn.run_sync(Base.metadata.create_all)
    except Exception as e:
        print(f"Database initialization error: {e}")
    
    # Start cleanup task
    cleanup_task = asyncio.create_task(periodic_cleanup())
    
    yield
    
    # Cleanup on shutdown
    cleanup_task.cancel()
    rag_sessions.clear()

async def periodic_cleanup():
    """Periodically clean up expired sessions"""
    while True:
        try:
            await asyncio.sleep(3600)  # Run every hour
            await cleanup_expired_sessions()
        except asyncio.CancelledError:
            break
        except Exception as e:
            print(f"Error in periodic cleanup: {e}")

app = FastAPI(lifespan=lifespan)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
async def root():
    return {"message": "FastAPI Login Server with Stateless RAG is running"}

class LoginRequest(BaseModel):
    username: str
    password: str

@app.post("/register")
async def register_user(request: LoginRequest):
    username = request.username
    password = request.password
    
    async with AsyncSessionLocal() as db:
        user = await get_user_by_username(db, username)
        if not user:
            # User doesn't exist, create new user
            new_user = await create_user(db, username, password)
            if new_user:
                return {
                    "message": "User registered successfully",
                    "username": username,
                    "status": "success"
                }
            else:
                raise HTTPException(
                    status_code=500,
                    detail="Failed to create user. Please try again."
                )
        else:
            # User already exists
            raise HTTPException(
                status_code=409,
                detail="Username already in use"
            )

SLOT_ASSIGNMENTS: Dict[str, int] = {}  # username -> slot mapping
AVAILABLE_SLOTS: Set[int] = set()  # available slot numbers
NEXT_SLOT: int = 0  # next slot number to assign if no available slots
SLOT_LOCK = asyncio.Lock()  # thread safety for slot operations

class LoginResponse(BaseModel):
    message: str
    username: str

@app.post("/login", response_model=LoginResponse)
async def login(request: LoginRequest):
    """Authenticate user with username and password"""
    username = request.username.strip()
    password = request.password
    
    async with AsyncSessionLocal() as db:
        user = await authenticate_user(db, username, password)
        if not user:
            raise HTTPException(
                status_code=401,
                detail="Invalid username or password"
            )
    
    return LoginResponse(
        message="Login successful",
        username=username,
    )

class LogoutRequest(BaseModel):
    username: str

@app.post("/logout")
async def logout(request: LogoutRequest):      
    return {"message": f"Logged out successfully."}

class MessageRequest(BaseModel):
    chatHistory: str
    message: str
    model : str
    ragStatus: bool
    session_id: str
    user : str

class MessageResponse(BaseModel):
    message: str

async def gen_context(session_id, query):
    """Query the RAG system for a specific session"""
    start_time = time.time()
    
    if session_id not in rag_sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    session = rag_sessions[session_id]
    session.update_access_time()
    
    if not session.has_document():
        raise HTTPException(
            status_code=400,
            detail="No document found in session. Please upload a document first."
        )
    
    if not query:
        raise HTTPException(status_code=400, detail="Query cannot be empty")
    
    try:
        # Get query embedding
        print(f"Getting embedding for query: {query}")
        query_embedding = await get_embeddings([query])
        print(f"Query embedding shape: {query_embedding.shape}")
        
        # Ensure embedding is float32 and normalized
        query_embedding = query_embedding.astype(np.float32)
        faiss.normalize_L2(query_embedding)
        
        # Validate vector store
        if session.vector_store is None:
            raise HTTPException(status_code=500, detail="Vector store not initialized")
        
        # Search similar chunks
        k = min(4, len(session.chunks))
        print(f"Searching for {k} similar chunks")
        
        # Perform the search
        scores, indices = session.vector_store.search(query_embedding, k)
        print(f"Search completed. Scores: {scores}, Indices: {indices}")
        
        # Get relevant chunks
        chunk_processing_start = time.time()
        relevant_chunks = []
        sources = set()
        for i, idx in enumerate(indices[0]):
            if 0 <= idx < len(session.chunks):  # Validate index bounds
                chunk = session.chunks[idx]
                relevant_chunks.append(chunk.page_content)
                sources.add(chunk.metadata.get("source", "Unknown"))
            else:
                print(f"Warning: Invalid chunk index {idx}")
        
        if not relevant_chunks:
            raise HTTPException(status_code=500, detail="No relevant chunks found")
        
        # Create context for LLM
        context = "\n\n".join([f"chunk{i+1}:\n{chunk}" for i, chunk in enumerate(relevant_chunks)])
        print(f"Context created with {len(relevant_chunks)} chunks")
        
        return context
        
    except Exception as e:
        total_time = time.time() - start_time
        print(f"Error occurred after {total_time:.2f} seconds: {e}")
        return "empty"

def llama_template(messages):
    """
    Convert OpenAI-style messages to Llama 3.1 chat template format.
    
    Llama 3.1 uses this format:
    <|begin_of_text|><|start_header_id|>system<|end_header_id|>
    
    {system_message}<|eot_id|><|start_header_id|>user<|end_header_id|>
    
    {user_message}<|eot_id|><|start_header_id|>assistant<|end_header_id|>
    
    {assistant_message}<|eot_id|>
    
    Args:
        messages (list): List of message dictionaries with 'role' and 'content' keys
    
    Returns:
        str: Formatted prompt string for Llama 3.1
    """
    if not messages:
        return "<|begin_of_text|>"
    
    prompt = "<|begin_of_text|>"
    
    for message in messages:
        role = message.get('role', '').lower().strip()
        content = message.get('content', '').strip()
        
        if not content:
            continue
        
        # Handle different roles
        if role in ['system', 'user', 'assistant']:
            prompt += f"<|start_header_id|>{role}<|end_header_id|>\n\n{content}<|eot_id|>"
        else:
            # Treat unknown roles as user messages
            prompt += f"<|start_header_id|>user<|end_header_id|>\n\n{content}<|eot_id|>"
    
    # If the last message is not from assistant, add assistant header for continuation
    if messages and messages[-1].get('role', '').lower() != 'assistant':
        prompt += "<|start_header_id|>assistant<|end_header_id|>\n\n"
    
    return prompt

@app.post("/message/stream")
async def send_message_stream(request: MessageRequest):
    """
    Send a message to the current llama.cpp model with streaming response
    """
    overall_start = time.time()
    
    history = request.chatHistory
    prompt = request.message.strip()
    ragStatus = request.ragStatus
    session_id= request.session_id
    port=AVAILABLE_PORT
    
    if not prompt.strip():
        raise HTTPException(status_code=400, detail="Message cannot be empty")
    
    print(f"Current model: {current_model}")
    print(f'Rag status : {ragStatus}')
    
    # Generating Chat History
    messages = []
    
    # Add system messages
    messages.append({
        "role": "system",
        "content": "You are a helpful AI assistant 'Varuna'. Provide accurate, concise, and engaging responses. Be conversational and friendly while staying professional. Give direct answers with relevant context. Acknowledge uncertainty rather than guessing. Ask clarifying questions when needed."
    })

    if history and history.strip():
        #messages seperate by \n
        lines = history.strip().split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('V: '):
                messages.append({
                    "role": "assistant",
                    "content": line[3:]  # Remove 'V: '
                })
            elif line.startswith('User: '):
                messages.append({
                    "role": "user",
                    "content": line[6:]  # Remove 'User: '
                })
    
    
    # RAG Context Generation
    if ragStatus:
        messages = []
        messages.append({
            "role": "system",
            "content": "You are a helpful AI assistant 'Varuna'. Provide accurate, concise, and engaging responses. Be conversational and friendly while staying professional. Give direct answers with relevant context. Acknowledge uncertainty rather than guessing. Ask clarifying questions when needed."
        })
        context = await gen_context(session_id,prompt)
        messages.append({
    "role": "system",
    "content": f"""DOCUMENT CONTEXT:
The user has uploaded a file. Below are up to 4 of the most relevant extracted content chunks from the document or image (OCR extracted text, if applicable). These have been selected because they are likely related to the user's question:

{context}

INSTRUCTIONS:
- Use only the information provided in the context above to answer the user's question.
- If the user has uploaded an image, the context will include extracted text from that image. Treat it the same as a text-based document.
- If the context is insufficient to answer the question fully, **clearly state what information is missing or unclear**.
- Do not fabricate details or assume anything beyond the provided context.
- If the user asks about an image, do not worry. The above context contains relevant information regarding the image 
ANSWER:
Provide a clear, concise, and comprehensive response based strictly on the context above.
"""
})

    # Add current user message
    messages.append({
        "role": "user",
        "content": prompt
    })

    # Format prompt
    formatted_prompt= llama_template(messages)
    format_time = time.time() - overall_start
    print(f"Prompt took: {format_time:.2f} seconds")
    
    print(messages)
    
    async def generate_stream():
        stream_start = time.time()
        first_token_time = None
        token_count = 0
        
        try:
            # Check if llama.cpp server is running
            client_start = time.time()
            async with httpx.AsyncClient(timeout=18000.0) as client:
                client_setup_time = time.time() - client_start
                print(f"HTTP client setup took: {client_setup_time:.4f} seconds")
                
                # Open API standard call with streaming
                print(f"Sending streaming request to llama.cpp server")
                
                request_start = time.time()
                async with client.stream(
                    method="POST",
                    url=f"{MODEL_PATH}{port}/completion",
                    json={
                        "model": current_model,
                        "prompt": formatted_prompt,
                        "temperature": 0.3,
                        "top_p": 0.95,
                        "max_tokens": 1000,
                        "stream": True,
                        "stop": ["<|eot_id|>", "<|end_of_text|>"],
                    },
                    headers={
                        "Content-Type": "application/json"
                    },
                    timeout=18000.0
                ) as response:
                    
                    request_send_time = time.time() - request_start
                    print(f"Request sending took: {request_send_time:.4f} seconds")
                    
                    if response.status_code != 200:
                        print(f"llama.cpp API error: {response.status_code}")
                        yield f"data: {json.dumps({'error': f'llama.cpp API returned status {response.status_code}'})}\n\n"
                        return
                    
                    # Process streaming response
                    chunk_processing_start = time.time()
                    async for chunk in response.aiter_lines():
                        if chunk:
                            # Remove 'data: ' prefix if present
                            if chunk.startswith('data: '):
                                chunk = chunk[6:]
                            
                            # Skip empty lines and [DONE] marker
                            if not chunk.strip() or chunk.strip() == '[DONE]':
                                continue
                            
                            try:
                                # Parse the JSON chunk
                                data = json.loads(chunk)
                                
                                if "content" in data:
                                    content = data["content"]
                                    if content:
                                        # Track first token time
                                        if first_token_time is None:
                                            first_token_time = time.time() - stream_start
                                            print(f"Time to first token: {first_token_time:.4f} seconds")
                                        
                                        token_count += 1
                                        # Send the token to the frontend
                                        yield f"data: {json.dumps({'token': content})}\n\n"

                                # Check if streaming is finished
                                if data.get("stop", False):
                                    total_stream_time = time.time() - stream_start
                                    tokens_per_second = token_count / total_stream_time if total_stream_time > 0 else 0
                                    print(f"Streaming completed - Total tokens: {token_count}, Total time: {total_stream_time:.4f}s, Tokens/sec: {tokens_per_second:.2f}")
                                    yield f"data: {json.dumps({'done': True})}\n\n"
                                        
                            except json.JSONDecodeError as e:
                                print(f"JSON decode error: {e}, chunk: {chunk}")
                                continue
                            except Exception as e:
                                print(f"Error processing chunk: {e}")
                                continue
                
        except HTTPException:
            error_time = time.time() - stream_start
            print(f"HTTP exception occurred after {error_time:.4f} seconds")
            yield f"data: {json.dumps({'error': 'HTTP exception occurred'})}\n\n"
        except Exception as e:
            error_time = time.time() - stream_start
            print(f"llama.cpp Streaming Error after {error_time:.4f} seconds: {str(e)}")
            yield f"data: {json.dumps({'error': f'Error generating response: {str(e)}'})}\n\n"
    
    return StreamingResponse(
        generate_stream(),
        media_type="text/plain",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Content-Type": "text/event-stream",
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Headers": "Content-Type",
        }
    )

# Document processing functions
def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file using PyMuPDF"""
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text() + "\n"
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    doc = DocxDocument(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file"""
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def extract_tables_from_pdf(file_path: str) -> List[Dict[str, Any]]:
    """Extract tables from PDF using pdfplumber"""
    tables_data = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                tables = page.extract_tables()
                for table_num, table in enumerate(tables):
                    if table and len(table) > 0:
                        headers = table[0] if table else [f"Column_{i}" for i in range(len(table[0]))]
                        rows = table[1:] if len(table) > 1 else []
                        tables_data.append({
                            'page': page_num + 1,
                            'table_number': table_num + 1,
                            'headers': [str(h) for h in headers if h],
                            'rows': [[str(cell) for cell in row] for row in rows if any(cell for cell in row)],
                            'type': 'table'
                        })
        print(f"Extracted {len(tables_data)} tables from PDF")
        return tables_data
    except Exception as e:
        print(f"Error extracting tables from PDF: {e}")
        raise

def extract_tables_from_docx(file_path: str) -> List[Dict[str, Any]]:
    """Extract tables from DOCX file"""
    tables_data = []
    try:
        doc = DocxDocument(file_path)
        for table_num, table in enumerate(doc.tables):
            if len(table.rows) > 0:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                rows = [[cell.text.strip() for cell in row.cells] for row in table.rows[1:] if any(cell.text.strip() for cell in row.cells)]
                tables_data.append({
                    'table_number': table_num + 1,
                    'headers': headers,
                    'rows': rows,
                    'type': 'table'
                })
        print(f"Extracted {len(tables_data)} tables from DOCX")
        return tables_data
    except Exception as e:
        print(f"Error extracting tables from DOCX: {e}")
        raise

def table_to_text(table_data: Dict[str, Any]) -> str:
    """Convert table data to readable text format"""
    if not table_data.get('headers') or not table_data.get('rows'):
        return ""

    text_lines = []
    table_id = f"TABLE {table_data['table_number']}"
    if 'page' in table_data:
        table_id += f" - PAGE {table_data['page']}"

    text_lines.append(f"[{table_id}]")
    text_lines.append(" | ".join(str(h) for h in table_data['headers']))
    text_lines.append("-" * len(" | ".join(str(h) for h in table_data['headers'])))

    for row in table_data['rows']:
        if len(row) == len(table_data['headers']):
            text_lines.append(" | ".join(str(cell) for cell in row))

    text_lines.append("[END TABLE]\n")
    return "\n".join(text_lines)

def table_to_structured_text(table_data: Dict[str, Any]) -> str:
    """Convert table to structured question-answer format for better RAG"""
    if not table_data.get('headers') or not table_data.get('rows'):
        return ""

    structured_text = []
    table_id = f"Table {table_data['table_number']}"
    if 'page' in table_data:
        table_id += f" on page {table_data['page']}"

    structured_text.append(f"=== {table_id} ===")

    for row_idx, row in enumerate(table_data['rows']):
        if len(row) == len(table_data['headers']):
            row_text = f"Row {row_idx + 1}: "
            field_descriptions = [f"{header}: {value}" for header, value in zip(table_data['headers'], row) if value and str(value).strip()]
            row_text += "; ".join(field_descriptions)
            structured_text.append(row_text)

    structured_text.append("=== End Table ===\n")
    return "\n".join(structured_text)

def extract_text_from_pdf_with_tables(file_path: str) -> Tuple[str, List[Dict[str, Any]]]:
    """Extract both text and tables from PDF"""
    text = extract_text_from_pdf(file_path)
    tables = extract_tables_from_pdf(file_path)
    return text, tables

def extract_text_from_docx_with_tables(file_path: str) -> Tuple[str, List[Dict[str, Any]]]:
    """Extract both text and tables from DOCX"""
    text = extract_text_from_docx(file_path)
    tables = extract_tables_from_docx(file_path)
    return text, tables

def create_chunks_with_tables(text: str, tables: List[Dict[str, Any]], filename: str) -> Tuple[List[Document], List[Document]]:
    """Split text into chunks and create separate table chunks with improved parameters"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=MIN_CHUNK_OVERLAP,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""]
    )

    text_chunks = text_splitter.split_text(text)
    text_documents = [
        Document(
            page_content=chunk,
            metadata={
                "source": filename,
                "chunk_id": i,
                "total_chunks": len(text_chunks),
                "type": "text"
            }
        ) for i, chunk in enumerate(text_chunks)
    ]

    table_documents = []
    for i, table in enumerate(tables):
        readable_table = table_to_text(table)
        structured_table = table_to_structured_text(table)
        combined_content = f"{structured_table}\n\n{readable_table}"

        doc = Document(
            page_content=combined_content,
            metadata={
                "source": filename,
                "chunk_id": f"table_{i}",
                "table_number": table.get('table_number', i + 1),
                "page": table.get('page', 'unknown'),
                "type": "table",
                "headers": table.get('headers', []),
                "row_count": len(table.get('rows', []))
            }
        )
        table_documents.append(doc)

    print(f"Created {len(text_documents)} text chunks and {len(table_documents)} table chunks")
    return text_documents, table_documents

def extract_text_from_img(file_path:str) -> str:
    try:
        # Read text from image
        results = ocr_reader.readtext(file_path)
        
        # Extract text from results
        text = ' '.join([result[1] for result in results])
        
        return text
        
    except Exception as e:
        print("Error during OCR:", e)
        return ""

def generate_caption(image_path):

    image = Image.open(image_path).convert("RGB")
    inputs = caption_processor(image, return_tensors="pt")

    out_ids = caption_model.generate(**inputs)
    caption = caption_processor.decode(out_ids[0], skip_special_tokens=True)
    return caption

def process_pdf_images(file_path: str) -> str:
    import fitz 
    
    result = ""
    temp_image_files = []  # Keep track of temporary files for cleanup
    image_counter = 1  # Counter for sequential image naming
    
    # Open the PDF document
    doc = fitz.open(file_path)
    
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        
        # Get list of images on the page
        image_list = page.get_images(full=True)
        
        for idx, img in enumerate(image_list):
            # Get the image reference
            xref = img[0]
            
            # Extract image data
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            
            # Create image name and path
            image_name = f"image{image_counter:02d}.png"
            image_path = f"/tmp/{image_name}"
            temp_image_files.append(image_path)
            
            try:
                # Load image from bytes and save as PNG
                pil_image = Image.open(io.BytesIO(image_bytes))
                
                # Convert to RGB if necessary
                if pil_image.mode in ('RGBA', 'LA'):
                    pil_image.save(image_path, "PNG")
                else:
                    pil_image = pil_image.convert('RGB')
                    pil_image.save(image_path, "PNG")
                
                result += f"Image description \n Page:{page_num} Image Number: {idx}\n"
                ocr_result = extract_text_from_img(image_path)
                result += f"Text extracted from image:\n {ocr_result}\n"
                caption = generate_caption(image_path)
                result += f"Image Caption:\n {caption}\n"
                
                image_counter += 1
                
            except Exception as e:
                print(f"Warning: Could not process image from page {page_num}: {e}")
                # Remove the path from temp_image_files if extraction failed
                if image_path in temp_image_files:
                    temp_image_files.remove(image_path)
    
    # Close the document
    doc.close()
    
    # Delete extracted images
    for temp_file in temp_image_files:
        if os.path.exists(temp_file):
            os.remove(temp_file)
    
    return result

async def get_embeddings(texts: List[str]) -> np.ndarray:
    """Get embeddings from llama.cpp server with Nomic model"""
    # Prepend required prefix to all texts
    prefixed_texts = [f"search_query: {text}" for text in texts]
    
    async with httpx.AsyncClient(timeout=60.0) as client:
        try:
            response = await client.post(
                url=f"{EMBEDDING_PATH}/v1/embeddings",
                json={"input": prefixed_texts},  # Send all texts in one request
                timeout=500.0
            )
            
            if response.status_code == 200:
                data = response.json()
                # Extract embeddings from all items in response data
                return np.array([item["embedding"] for item in data["data"]], dtype=np.float32)
            else:
                raise Exception(f"API error {response.status_code}: {response.text}")
                
        except Exception as e:
            print(f"Embedding error: {e}")
            raise

def create_chunks(text: str, filename: str) -> List[Document]:
    """Split text into chunks"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1024,
        chunk_overlap=20,
        length_function=len,
        separators=["\n\n", "\n", " ", ""]
    )
    
    chunks = text_splitter.split_text(text)
    documents = []
    
    for i, chunk in enumerate(chunks):
        doc = Document(
            page_content=chunk,
            metadata={
                "source": filename,
                "chunk_id": i,
                "total_chunks": len(chunks)
            }
        )
        documents.append(doc)
    
    return documents

# RAG Session endpoints
class SessionResponse(BaseModel):
    session_id: str
    status: str
    message: str

@app.post("/rag/create-session", response_model=SessionResponse)
async def create_rag_session():
    """Create a new RAG session"""
    session_id = str(uuid.uuid4())
    rag_sessions[session_id] = RAGSession(session_id)
    
    return SessionResponse(
        session_id=session_id,
        status="success",
        message="RAG session created successfully"
    )

@app.delete("/rag/session/{session_id}")
async def delete_rag_session(session_id: str):
    """Delete a RAG session"""
    if session_id not in rag_sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    del rag_sessions[session_id]
    
    return JSONResponse({
        "status": "success",
        "message": "Session deleted successfully"
    })

class RAGUploadResponse(BaseModel):
    status: str
    message: str
    chunks_created: int
    processing_time: float
    filename: str
    session_id: str

@app.post("/rag/upload/{session_id}", response_model=RAGUploadResponse)
async def upload_document_to_session(
    session_id: str,
    file: UploadFile = File(...),
    include_tables: bool = Form(True)
):
    """Upload and process a document for RAG with table extraction"""
    if session_id not in rag_sessions:
        print(f"Attempt to upload to non-existent session: {session_id}")
        raise HTTPException(status_code=404, detail="Session not found")

    session = rag_sessions[session_id]
    session.update_access_time()
    if not file.filename:
        print("No file provided for upload")
        raise HTTPException(status_code=400, detail="No file provided")

    if session.has_document():
        print(f"Session {session_id} already has document: {session.document_name}")
        raise HTTPException(status_code=400, detail=f"Session already has a document: {session.document_name}")

    file_extension = os.path.splitext(file.filename)[1].lower()
    supported_extensions = ['.pdf', '.txt', '.docx', '.doc','.png','.jpg','.jpeg']

    if file_extension not in supported_extensions:
        print(f"Unsupported file type: {file_extension}")
        raise HTTPException(status_code=400, detail=f"Unsupported file type. Supported types: {', '.join(supported_extensions)}")

    start_time = time.time()
    temp_file_path = None

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_file_path = temp_file.name

        file_size = len(content)
        file_size_mb = file_size / (1024 * 1024)
        #5MB max size
        skip_image_extraction = file_size > 5 * 1024 * 1024

        tables = []
        if file_extension == '.pdf':
            text, tables = extract_text_from_pdf_with_tables(temp_file_path) if include_tables else (extract_text_from_pdf(temp_file_path), [])
            # Only extract images if file is smaller than 50MB
            if not skip_image_extraction:
                image_data = process_pdf_images(temp_file_path)
                text += f"\n{image_data}"
            else:
                print(f"Skipping image extraction for large PDF ({file_size_mb:.2f}MB): {file.filename}")
                
        elif file_extension in ['.docx', '.doc']:
            text, tables = extract_text_from_docx_with_tables(temp_file_path) if include_tables else (extract_text_from_docx(temp_file_path), [])
        elif file_extension == '.txt':
            text = extract_text_from_txt(temp_file_path)
        elif file_extension == '.png' or file_extension=='.jpg' or file_extension == ".jpeg":
            text="The following information is the image extracted text and the image caption\n"
            text += extract_text_from_img(temp_file_path)
            text+="Image caption \n"
            text+=generate_caption(temp_file_path)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")

        if not text.strip():
            print(f"No text content found in document: {file.filename}")
            raise HTTPException(status_code=400, detail="No text content found in the document")

        if include_tables and tables:
            text_chunks, table_chunks = create_chunks_with_tables(text, tables, file.filename)
            all_chunks = text_chunks + table_chunks
        else:
            all_chunks = create_chunks(text, file.filename)
            table_chunks = []

        if not all_chunks:
            print(f"Failed to create chunks from document: {file.filename}")
            raise HTTPException(status_code=400, detail="Failed to create chunks from document")

        chunk_texts = [doc.page_content for doc in all_chunks]
        embeddings = await get_embeddings(chunk_texts)

        dimension = embeddings.shape[1]
        index = faiss.IndexFlatIP(dimension)
        faiss.normalize_L2(embeddings)
        index.add(embeddings)

        session.document_name = file.filename
        session.chunks = all_chunks
        session.embeddings = embeddings
        session.vector_store = index
        session.chunk_texts = chunk_texts
        session.tables = tables
        session.table_chunks = table_chunks

        processing_time = time.time() - start_time
        
        # Add file size info to log message
        log_message = f"Processed document {file.filename} ({file_size_mb:.2f}MB) with {len(all_chunks)} chunks in {processing_time:.2f} seconds"
        if skip_image_extraction:
            log_message += " (image extraction skipped due to file size)"
        print(log_message)

        response_message = f"Document processed successfully. Created {len(all_chunks)} chunks ({len(table_chunks)} table chunks, {len(tables)} tables extracted)."
        if skip_image_extraction:
            response_message += " Image extraction was skipped due to large file size (>50MB)."

        return RAGUploadResponse(
            status="success",
            message=response_message,
            chunks_created=len(all_chunks),
            processing_time=round(processing_time, 2),
            filename=file.filename,
            session_id=session_id
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error processing document {file.filename}: {e}")
        session.clear_document_data()
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
                print(f"Deleted temporary file: {temp_file_path}")
            except Exception as e:
                print(f"Error deleting temporary file {temp_file_path}: {e}")
                               
@app.delete("/rag/document/{session_id}")
async def remove_document_from_session(session_id: str):
    """Remove document and all processed data from session"""
    if session_id not in rag_sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    session = rag_sessions[session_id]
    session.update_access_time()
    
    if not session.has_document():
        raise HTTPException(status_code=400, detail="No document found in session")
    
    document_name = session.document_name
    session.clear_document_data()
    
    return JSONResponse({
        "status": "success",
        "message": f"Document '{document_name}' and all processed data removed from session"
    })

class RAGQueryRequest(BaseModel):
    query: str
    max_results: Optional[int] = 4

class RAGQueryResponse(BaseModel):
    response: str
    sources: List[str]
    relevant_chunks: List[str]
    status: str
    session_id: str

@app.get("/rag/session/{session_id}/status")
async def get_session_status(session_id: str):
    """Get status of a RAG session"""
    if session_id not in rag_sessions:
        print(f"Attempt to access status of non-existent session: {session_id}")
        raise HTTPException(status_code=404, detail="Session not found")

    session = rag_sessions[session_id]
    session.update_access_time()

    return JSONResponse({
        "session_id": session_id,
        "has_document": session.has_document(),
        "document_name": session.document_name,
        "chunks_count": len(session.chunks),
        "table_chunks_count": len(getattr(session, 'table_chunks', [])),
        "tables_extracted": len(getattr(session, 'tables', [])),
        "created_at": session.created_at.isoformat(),
        "last_accessed": session.last_accessed.isoformat(),
        "embedding_model": embedding_model,
        "status": "success"
    })

@app.get("/rag/session/{session_id}/tables")
async def get_session_tables(session_id: str):
    """Get information about tables in the session"""
    if session_id not in rag_sessions:
        print(f"Attempt to access tables of non-existent session: {session_id}")
        raise HTTPException(status_code=404, detail="Session not found")

    session = rag_sessions[session_id]
    session.update_access_time()

    if not hasattr(session, 'tables') or not session.tables:
        return JSONResponse({
            "session_id": session_id,
            "tables_count": 0,
            "tables": [],
            "status": "success"
        })

    table_summaries = [
        {
            "table_number": table.get('table_number', i + 1),
            "page": table.get('page', 'unknown'),
            "headers": table.get('headers', []),
            "row_count": len(table.get('rows', [])),
            "column_count": len(table.get('headers', []))
        } for i, table in enumerate(session.tables)
    ]

    print(f"Retrieved {len(table_summaries)} tables for session {session_id}")
    return JSONResponse({
        "session_id": session_id,
        "tables_count": len(session.tables),
        "tables": table_summaries,
        "status": "success"
    })

@app.get("/rag/sessions")
async def list_rag_sessions():
    """List all active RAG sessions"""
    await cleanup_expired_sessions()
    
    sessions_info = []
    for session_id, session in rag_sessions.items():
        sessions_info.append({
            "session_id": session_id,
            "has_document": session.has_document(),
            "document_name": session.document_name,
            "chunks_count": len(session.chunks),
            "created_at": session.created_at.isoformat(),
            "last_accessed": session.last_accessed.isoformat()
        })
    
    return JSONResponse({
        "active_sessions": len(rag_sessions),
        "sessions": sessions_info,
        "status": "success"
    })

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

"""
for all servers:-
export LD_LIBRARY_PATH=build/ggml/src:build/src/:$LD_LIBRARY_PATH

for embedding server:-
echo 0 | sudo tee /proc/sys/kernel/yama/ptrace_scope

./build/bin/llama-server -m "/media/caio/New Volume/Office_llm/models/Meta-Llama-3.1-8B-Instruct-Q4_K_M.gguf" -np  -c 2048 --port 8081 --host 0.0.0.0 --chat-template llama3

./build/bin/llama-server -m "/media/caio/New Volume/Office_llm/models/nomic-embed-text-v1.Q4_K_M.gguf" -c 2048 --port 8083 --host 0.0.0.0 --embeddings
"""